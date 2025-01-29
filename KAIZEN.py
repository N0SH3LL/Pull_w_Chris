import src.SCC.scc_read
import src.SCC.scc_check
import src.SCC.scc_tables
from src.utils import file_operations
import argparse
import os
import json
import shutil
import sys
import subprocess
import re
import docx
from datetime import datetime
from src.Tenable import api_client
from src.Tenable import report_operations

def read_json(filename): #needed for opening progress.json
    with open(filename, 'r') as file:
        return json.load(file)

def write_not_gathered_file(output_filename='Not_gathered.txt'): # obsolete/deprecated
    progress_data = read_json('progress.json')
    all_bper_dict = progress_data.get('BPERs', {})
    all_doc_dict = progress_data.get('Documents', {})
    all_attestation_dict = progress_data.get('Attestations', {})
    
    not_gathered_info = {}

    for data_dict in [all_attestation_dict, all_bper_dict, all_doc_dict]:
        for key, entries in data_dict.items():
            for value in entries:
                # Check for 'Gathered' being blank or 'TLA' being True in all_bper_dict
                if value.get('Gathered') == '' or (value.get('TLA') == True and data_dict is all_bper_dict):
                    scc_name = value['SCC']
                    if scc_name not in not_gathered_info:
                        not_gathered_info[scc_name] = []
                    detail = value.get('Attestation num', '') or value.get('BPER name', '') or value.get('Doc name', '')
                    if detail:
                        not_gathered_info[scc_name].append(detail)

    # Write to the output file
    with open(output_filename, 'w') as output_file:
        for scc_name, details in sorted(not_gathered_info.items()):
            output_file.write(f"{scc_name}\n")
            for detail in sorted(details):
                output_file.write(f"\t{detail}\n")
            output_file.write("\n")

def update_dict(all_dict, new_entries): # check the dictionary to make sure any value being added is unique
    for key, value in new_entries.items():
        if key not in all_dict:
            all_dict[key] = []
        all_dict[key].append(value)

def replace_text_in_docx(doc_path, replacements): #supports filling in some of the templates
    doc = docx.Document(doc_path)

    # Replace in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        for key, val in replacements.items():
                            if key in run.text:
                                run.text = run.text.replace(key, val)

    doc.save(doc_path)

def create_directories(project_dir):
    progress_file_path = os.path.join(project_dir, 'progress.json')

    try:
        with open(progress_file_path, 'r') as file:
            progress_data = json.load(file)
            scc_dict = progress_data.get('SCC', {})
            attestation_dict = progress_data.get('Attestations', {})
            project_settings = progress_data.get('Program Settings', {})
            project_directory = project_settings.get('Project Directory', project_dir)
    except (IOError, json.JSONDecodeError) as e:
        print(f"Failed to read or parse the progress.json file: {e}")
        return

    print(f"Creating directories in {project_directory}")
    print(f"Number of SCCs to process: {len(scc_dict)}")

    for scc_path, details in scc_dict.items():
        try:
            scc_name = details['SCC']
            sanitized_scc_name = re.sub(r'_\d{2}$', '', scc_name).strip()
            main_dir_path = os.path.join(project_directory, sanitized_scc_name)

            print(f"Processing SCC: {sanitized_scc_name}")

            if not details.get('Directory built', False):
                os.makedirs(main_dir_path, exist_ok=True)
                print(f"Created main directory: {main_dir_path}")

                # Create standard subdirectories
                standard_subdirs = ["Exceptions and Deviations", "Supporting Documents"]
                
                # Check if there are attestations for this SCC
                has_attestations = any(att_info.get('SCC') == scc_name for att_list in attestation_dict.values() for att_info in (att_list if isinstance(att_list, list) else [att_list]))
                
                if has_attestations:
                    standard_subdirs.append("Attestations")
                    print(f"Attestations found for {scc_name}, creating Attestations directory")
                else:
                    print(f"No attestations found for {scc_name}, skipping Attestations directory")

                for subdir in standard_subdirs:
                    subdir_path = os.path.join(main_dir_path, subdir)
                    os.makedirs(subdir_path, exist_ok=True)
                    print(f"Created standard subdirectory: {subdir_path}")

                # Create evidence-specific subdirectories
                evidence_methods = details.get('Evidence Methods', [])
                manual_subdirs = set()

                for method in evidence_methods:
                    method = method.lower()
                    if "automated" in method:
                        automated_path = os.path.join(main_dir_path, "Automated")
                        os.makedirs(automated_path, exist_ok=True)
                        print(f"Created Automated directory: {automated_path}")
                    
                    if "manual" in method:
                        manual_path = os.path.join(main_dir_path, "Manual")
                        os.makedirs(manual_path, exist_ok=True)
                        print(f"Created Manual directory: {manual_path}")

                        if "screenshot" in method:
                            manual_subdirs.add("Screenshots")
                        if "auto info" in method:
                            manual_subdirs.add("Automated Info")
                        if "script" in method:
                            manual_subdirs.add("Scripts")
                        if "document" in method:
                            manual_subdirs.add("Documents")
                        if "3rd party tool" in method:
                            manual_subdirs.add("3rd party tools")

                # Create Manual subdirectories
                for subdir in manual_subdirs:
                    subdir_path = os.path.join(main_dir_path, "Manual", subdir)
                    os.makedirs(subdir_path, exist_ok=True)
                    print(f"Created Manual subdirectory: {subdir_path}")

                details['Directory built'] = True
                print(f"Marked directory as built for SCC: {sanitized_scc_name}")
            else:
                print(f"Directory already built for SCC: {sanitized_scc_name}")

        except KeyError as e:
            print(f"Warning: Missing key {e} for SCC entry: {scc_path}")
            continue
        except Exception as e:
            print(f"Error creating directories for {scc_path}: {e}")
            continue

    # Save the updated progress data to progress.json
    try:
        with open(progress_file_path, 'w') as file:
            json.dump(progress_data, file, indent=4)
        print("Updated progress.json with directory build status")
    except IOError as e:
        print(f"Failed to write updated progress data: {e}")

    print("Directory creation process completed.")

def build_templates(method_dict, project_dir, template_dir):
    print("Building templates...")
    
    # Group STIG IDs by SCC
    scc_groups = {}
    for stig_id, details in method_dict.items():
        scc_name = details['SCC']
        sanitized_scc_name = re.sub(r'_\d{2}$', '', scc_name).strip()
        if sanitized_scc_name not in scc_groups:
            scc_groups[sanitized_scc_name] = []
        scc_groups[sanitized_scc_name].append((stig_id, details))

    for sanitized_scc_name, stig_items in scc_groups.items():
        print(f"Processing templates for SCC: {sanitized_scc_name}")
        main_dir_path = os.path.join(project_dir, sanitized_scc_name)
        
        # Check for different types of checks
        has_manual_document = any(details['Evidence method'].lower() == 'manual-document' for _, details in stig_items)
        has_automated = any('automated' in details['Evidence method'].lower() for _, details in stig_items)
        has_manual = any('manual' in details['Evidence method'].lower() for _, details in stig_items)
        manual_screenshot_stig_ids = [stig_id for stig_id, details in stig_items if details['Evidence method'].lower() == 'manual-screenshot']

        # Teamname-Document_Evidence
        if has_manual_document:
            manual_folder = os.path.join(main_dir_path, "Manual")
            if os.path.exists(manual_folder):
                source = os.path.join(template_dir, "Teamname-Document_Evidence.xlsx")
                dest = os.path.join(manual_folder, f"!{sanitized_scc_name}-Document_Evidence.xlsx")
                shutil.copy2(source, dest)
                print(f"Copied Document Evidence template to {dest}")

        # Teamname-EvidenceValidation
        if has_automated:
            automated_folder = os.path.join(main_dir_path, "Automated")
            if os.path.exists(automated_folder):
                source = os.path.join(template_dir, "Teamname-EvidenceValidation.xlsx")
                dest = os.path.join(automated_folder, f"!{sanitized_scc_name}-EvidenceValidation.xlsx")
                shutil.copy2(source, dest)
                print(f"Copied Evidence Validation template to {dest}")

        # Teamname-Manual_Control_Evidence
        if has_manual:
            manual_folder = os.path.join(main_dir_path, "Manual")
            if os.path.exists(manual_folder):
                source = os.path.join(template_dir, "Teamname-Manual_Control_Evidence.xlsx")
                dest = os.path.join(manual_folder, f"!{sanitized_scc_name}-Manual_Control_Evidence.xlsx")
                shutil.copy2(source, dest)
                print(f"Copied Manual Control Evidence template to {dest}")

        # Teamname-DeviceGapList and Teamname-Remediation
        source_gap = os.path.join(template_dir, "Teamname-DeviceGapList.xlsx")
        dest_gap = os.path.join(main_dir_path, f"!{sanitized_scc_name}-DeviceGapList.xlsx")
        shutil.copy2(source_gap, dest_gap)
        print(f"Copied Device Gap List template to {dest_gap}")

        source_rem = os.path.join(template_dir, "Teamname-Remediation.xlsx")
        dest_rem = os.path.join(main_dir_path, f"!{sanitized_scc_name}-Remediation.xlsx")
        shutil.copy2(source_rem, dest_rem)
        print(f"Copied Remediation template to {dest_rem}")

        # Manual Screenshot Template
        if manual_screenshot_stig_ids:
            screenshots_folder = os.path.join(main_dir_path, "Manual", "Screenshots")
            if os.path.exists(screenshots_folder):
                for stig_id in manual_screenshot_stig_ids:
                    source = os.path.join(template_dir, "Manual Screenshot Template.docx")
                    dest = os.path.join(screenshots_folder, f"{stig_id}.docx")
                    shutil.copy2(source, dest)
                    print(f"Copied Manual Screenshot template for STIG ID {stig_id} to {dest}")

                    # Replace placeholders in the copied file
                    replace_text_in_docx(dest, {
                        "FILENAMEINSERT": sanitized_scc_name,
                        "STIGIDINSERT": stig_id
                    })

    print("Template building completed.")

def convert_datetime_to_string(obj): # converts datetime values to string because it pitches a fit if not done
    if isinstance(obj, dict):
        return {key: convert_datetime_to_string(value) for key, value in obj.items()}
    elif isinstance(obj, list):
        return [convert_datetime_to_string(item) for item in obj]
    elif isinstance(obj, datetime):
        return obj.isoformat()
    return obj

def save_master_dicts(data, filename): # save the dictionaries after converting the values to string
    converted_data = convert_datetime_to_string(data)
    with open(filename, 'w') as file:
        json.dump(converted_data, file, indent=4)

# TODO remove below function and replace with a copy_sccs_to_folders
def move_sccs_to_folders(project_dir, file_extensions): # moves the scc's into their respetive directories ****NOT USED IN GUI, NEEDS TO COPY INSTEAD****
    for file in os.listdir(project_dir):
        file_path = os.path.join(project_dir, file)
        file_name, file_ext = os.path.splitext(file)

        # Apply the sanitization only for Excel files
        if file_ext in ['.xlsx', '.xls']:
            sanitized_folder_name = re.sub(r'_\d{2}$', '', file_name).strip()
        else:
            sanitized_folder_name = file_name

        if file_ext in file_extensions:
            dest_folder_path = os.path.join(project_dir, sanitized_folder_name)

            if os.path.exists(dest_folder_path) and os.path.isdir(dest_folder_path):
                shutil.move(file_path, os.path.join(dest_folder_path, file))
            else:
                print(f"Destination folder not found for {file}, expected at {dest_folder_path}")

def build_progress_json(directory_path, project_dir):
    #Master dictionaries
    all_bper_dict = {}
    all_doc_dict = {}
    all_attestation_dict = {}
    scc_data_dict = {}
    checks_data_dict = {}

    #Loop for all excel files in directory
    for file in os.listdir(directory_path):
        if file.endswith('.xlsx') or file.endswith('.xls'):
            file_path = os.path.join(directory_path, file)

            #This creates the dictionaries using src.SCC.scc_read (bper, docs, atts, controls)
            bper_dict, doc_dict, attestation_dict, method_dict = src.SCC.scc_read.process_excel_file(file_path)

            #This creates the scc_info dictionary (first page check, etc)
            scc_info = src.SCC.scc_check.process_scc_file(file_path)

            #stores in master dict
            update_dict(all_bper_dict, bper_dict)
            update_dict(all_attestation_dict, attestation_dict)
            update_dict(all_doc_dict, doc_dict)

            # Update scc_data_dict and checks_data_dict
            scc_data_dict[file_path] = scc_info
            for stig_id, details in method_dict.items():
                # Extract SCC name from the file path and remove extension and trailing "_**"
                scc_name = os.path.splitext(os.path.basename(file_path))[0]
                scc_name = re.sub(r'_\d{2}$', '', scc_name).strip()
                checks_data_dict[stig_id] = {
                    'SCC': scc_name,
                    'Evidence method': details['Evidence Method']
                }
           # print(f"Added {len(method_dict)} checks for SCC: {scc_name}")
            scc_methods = set()
            for check_info in checks_data_dict.values():
                if check_info['SCC'] == scc_name:
                    scc_methods.add(check_info['Evidence method'].lower())
            scc_info['Evidence Methods'] = list(scc_methods)
            
    # Save progress to progress.json
    progress_data = {
        'BPERs': all_bper_dict,
        'Attestations': all_attestation_dict,
        'Documents': all_doc_dict,
        'SCC': scc_data_dict,
        'Checks': checks_data_dict,
        'Program Settings': {
            'Project Directory': project_dir,
            'Directories Built': False,
            'Templates Built': False,
            'Gather and Sort Date': '',
            'Doc Tracker Update': '',
            'Pull Info Date': '',
            'Checklists generated':''
        }
    }
    with open(os.path.join(project_dir, 'progress.json'), 'w') as file:
        json.dump(convert_datetime_to_string(progress_data), file, indent=4)
    #print(f"Total checks in progress.json: {len(checks_data_dict)}")

def update_bper_dict(master_directory):
    progress_data = read_json('progress.json')
    bper_dict = progress_data.get('BPERs', {})

    for key, value_list in bper_dict.items():
        for value in value_list:
            scc_name = value['SCC']
            file_path = os.path.join(master_directory, scc_name, "Exceptions and Deviations", f"{key}.pdf")
            if os.path.isfile(file_path):
                valid_to_date, approval_status, tla_present = file_operations.extract_BPER_info(file_path)
                value['Valid to'] = valid_to_date
                value['Approval Status'] = approval_status
                value['TLA'] = tla_present
                print(f"Updated BPER '{key}' - 'Valid to': {valid_to_date}, 'Approval Status': {approval_status}, 'TLA': {tla_present}")
            else:
                print(f"File not found for BPER: {key}")

    progress_data['BPERs'] = bper_dict
    with open('progress.json', 'w') as file:
        json.dump(convert_datetime_to_string(progress_data), file, indent=4)

def update_attestation_dict(master_directory):
    progress_data = read_json('progress.json')
    attestation_dict = progress_data.get('Attestations', {})

    for key, value_list in attestation_dict.items():
        for value in value_list:
            scc_name = value['SCC']
            file_path = os.path.join(master_directory, scc_name, "Attestations", f"{key}.pdf")
            if os.path.isfile(file_path):
                approval_status, approval_date = file_operations.extract_attest_info(file_path)
                value['Approval Status'] = approval_status
                value['Valid to'] = approval_date
                print(f"Updated '{key}' with status: {approval_status} and date: {approval_date}")
            else:
                print(f"File not found for Attestation: {key}")

    progress_data['Attestations'] = attestation_dict
    with open('progress.json', 'w') as file:
        json.dump(convert_datetime_to_string(progress_data), file, indent=4)

def update_doc_dict(master_directory):
    progress_data = read_json('progress.json')
    doc_dict = progress_data.get('Documents', {})

    for key, value_list in doc_dict.items():
        for value in value_list:
            scc_name = value['SCC']
            file_path = os.path.join(master_directory, scc_name, "Supporting Documents", f"{key}.docx")
            if os.path.isfile(file_path):
                most_recent_date = file_operations.extract_Doc_info(file_path)
                if most_recent_date:
                    value['Last update'] = most_recent_date
                    print(f"Updated '{key}' with 'Last update': {most_recent_date}")
            else:
                print(f"File not found for Document: {key}")

    progress_data['Documents'] = doc_dict
    with open('progress.json', 'w') as file:
        json.dump(convert_datetime_to_string(progress_data), file, indent=4)

def gather_and_process_reports(project_dir):
    client = api_client()
    owner_id = "0029702"  # My hardcoded ID !!!REPLACE WITH YOURS!!! This is only used here to filter on the reports

    try:
        # Download reports
        report_operations.download_reports_for_owner(client, owner_id, project_dir)
        print("Reports downloaded successfully.")

        # Process reports
        for scc_name in os.listdir(project_dir):
            scc_dir = os.path.join(project_dir, scc_name)
            if os.path.isdir(scc_dir):
                automated_dir = os.path.join(scc_dir, "Automated")
                if os.path.exists(automated_dir):
                    print(f"Processing reports for {scc_name}")
                    run_powershell_script(automated_dir)
                    organize_output_files(automated_dir)

        print("All reports processed and organized.")
    except Exception as e:
        print(f"Error in gather_and_process_reports: {str(e)}")
        raise

def run_powershell_script(folder_path):
    script_path = os.path.join(os.path.dirname(__file__), "scripts", "SCReport-Parse-Multiple 1.ps1")
    try:
        subprocess.run(["powershell", "-ExecutionPolicy", "Bypass", "-File", script_path], 
                       cwd=folder_path, check=True, capture_output=True, text=True)
        print(f"PowerShell script executed successfully in {folder_path}")
    except subprocess.CalledProcessError as e:
        print(f"Error executing PowerShell script in {folder_path}: {e}")
        print(f"Script output: {e.output}")

def organize_output_files(folder_path):
    # use existing tenable scripts, and put them in the correct folders
    metadata_folder = os.path.join(folder_path, "MetaData")
    failed_folder = os.path.join(folder_path, "Failed")
    
    os.makedirs(metadata_folder, exist_ok=True)
    os.makedirs(failed_folder, exist_ok=True)
    
    for file in os.listdir(folder_path):
        if file.endswith("_MetaData.csv"):
            os.rename(os.path.join(folder_path, file), os.path.join(metadata_folder, file))
            print(f"Moved {file} to MetaData folder")
        elif file.endswith("_FailedChecks.csv"):
            os.rename(os.path.join(folder_path, file), os.path.join(failed_folder, file))
            print(f"Moved {file} to Failed folder")

def main():
    parser = argparse.ArgumentParser(description='Main script to process SCC files.')
    parser.add_argument('directory_path', type=str, help='Path to the directory containing Excel files')
    parser.add_argument('--progress', action='store_true', help='Load progress from progress.json')
    args = parser.parse_args()

    if not os.path.isdir(args.directory_path):
        print(f"Directory not found: {args.directory_path}")
        return

    #Checks and splits bulk BPERs
    # All_BPER_path = os.path.join(args.directory_path, 'All BPERs')
    # split_bper.process_directory(All_BPER_path)
    # master_directory = args.directory_path
    
    # Load progress from progress.json if --progress flag is set
    if args.progress:
        progress_file = 'progress.json'
        if os.path.exists(progress_file):
            with open(progress_file, 'r') as file:
                progress_data = json.load(file)
        else:
            print(f"Progress file {progress_file} not found. Starting from scratch.")
            progress_data = {
                'BPERs': {},
                'Attestations': {},
                'Documents': {},
                'SCC': {},
                'Checks': {},
                'Program Settings': {
                    'Project Directory': master_directory,
                    'Directories Built': False,
                    'Templates Built': False,
                    'Gather and Sort Date': '',
                    'Doc Tracker Update': ''
                }
            }
    else:
        progress_data = {
            'BPERs': {},
            'Attestations': {},
            'Documents': {},
            'SCC': {},
            'Checks': {},
            'Program Settings': {
                'Project Directory': master_directory,
                'Directories Built': False,
                'Templates Built': False,
                'Gather and Sort Date': '',
                'Doc Tracker Update': ''
            }
        }

    #Loop for all excel files in directory
    for file in os.listdir(args.directory_path):
        if file.endswith('.xlsx') or file.endswith('.xls'):
            file_path = os.path.join(args.directory_path, file)

            # Check if the file has already been processed
            if file_path in progress_data['SCC']:
                print(f"Skipping already processed file: {file_path}")
                continue

            #This creates the dictionaries using src.SCC.scc_read (bper, docs, atts, controls)
            bper_dict, doc_dict, attestation_dict, method_dict = src.SCC.scc_read.process_excel_file(file_path)
            #This creates the scc_info dictionary (first page check, etc)
            scc_info = src.SCC.scc_check.process_scc_file(file_path)

            #This creates the directories for each SCC
            build_templates(master_directory)

            #This grabs the files and updates the dictionaries
            base_directories = {'bper': 'All BPERs', 'doc': 'All Docs', 'attestation': 'All Attestations'}
            file_operations.update_dictionaries_and_copy_files(bper_dict, doc_dict, attestation_dict, base_directories, master_directory)

            scc_name_without_extension = os.path.splitext(os.path.basename(file_path))[0]
            scc_name_without_extension = re.sub(r'_\d{2}$', '', scc_name_without_extension).strip()
            print(f"\nGathered documents for {scc_name_without_extension}")

            #This writes the checklist for each SCC
            src.SCC.scc_tables.write_checklist(bper_dict, doc_dict, attestation_dict, method_dict, scc_info, master_directory)

            #stores in master dict
            update_dict(progress_data['BPERs'], bper_dict)
            update_dict(progress_data['Attestations'], attestation_dict)
            update_dict(progress_data['Documents'], doc_dict)

            # Update scc_data_dict and checks_data_dict
            progress_data['SCC'][file_path] = scc_info
            for stig_id, details in method_dict.items():
                # Extract SCC name from the file path and remove extension and trailing "_**"
                scc_name = os.path.splitext(os.path.basename(file_path))[0]
                scc_name = re.sub(r'_\d{2}$', '', scc_name).strip()
                progress_data['Checks'][stig_id] = {
                    'SCC': scc_name,
                    'Evidence method': details['Method']
                }

    #move text and SCC's into their folders
    file_extensions = ['.xlsx', '.txt']
    move_sccs_to_folders(master_directory, file_extensions)

    #saves master dicts
    save_master_dicts(progress_data['Attestations'], 'all_attestation_dict.json')
    save_master_dicts(progress_data['BPERs'], 'all_bper_dict.json')
    save_master_dicts(progress_data['Documents'], 'all_doc_dict.json')

    #build not gathered list
    write_not_gathered_file()

    # Save progress to progress.json
    with open('progress.json', 'w') as file:
        json.dump(convert_datetime_to_string(progress_data), file, indent=4)


if __name__ == "__main__":
    main()
    print('done')
